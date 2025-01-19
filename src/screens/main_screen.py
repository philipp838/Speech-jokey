# Kivy
from kivy.uix.floatlayout import FloatLayout
from kivy.metrics import sp
from kivy.app import App
from kivy.clock import Clock
from kivy.properties import ObjectProperty, StringProperty
from kivy.logger import Logger as log
# KivyMD
from kivymd.uix.screen import MDScreen
from kivymd.uix.menu import MDDropdownMenu
from kivymd.uix.filemanager import MDFileManager
from kivymd.uix.dialog import MDDialog
from kivy.uix.popup import Popup
# stdlib
import os
import sys
# Custom
from modules.dialog.exitdialog import ExitDialog
# .docx to text conversion
import docx2txt
from docx import Document


class EmojiPopup(Popup):
    def __init__(self, text_input, **kwargs):
        super(EmojiPopup, self).__init__(**kwargs)
        self.text_input = text_input
        # Access the shared ssml tags dictionary from MainScreen
        self.ssml_tags = MainScreen.ssml_tags
        # Track the state for each emoji
        self.tag_state = {emoji: "open" for emoji in self.ssml_tags.keys()}

    # Old method to place emoji in textfield
    def insert_emoji(self, emoji):
        # Get cursor position
        cursor_index = self.text_input.cursor_index()

        # Divide the current text of the text input into two parts:
        # Text before and after cursor position
        current_text = self.text_input.text
        new_text = (
                current_text[:cursor_index] +  # Text in front of cursor
                emoji +  # The emoji to be inserted
                current_text[cursor_index:]  # Text after cursor
        )

        # Set the new text in the TextInput
        self.text_input.text = new_text

        # Place the cursor behind the inserted emoji
        self.text_input.cursor = (cursor_index + len(emoji), 0)

        self.dismiss()


from api.api_factory import api_factory


class MainScreen(MDScreen):
    title = StringProperty()
    current_engine_text = StringProperty("tts engine: \n")
    text_input = ObjectProperty(None)
    text_type = StringProperty("text")
    voice_dialog = None
    selected_voice = StringProperty()
    # FIXME The values of this dictionary needs to be kept in sync with the screen names in main.py (Unfortunately)
    menu_options = {
        "Settings": "settings",
        "About": "about",
        "Exit": None  # NOTE Exit just closes the app and doesn't have an associated screen
    }
    ssml_tags = {
        "‚è∏Ô∏è": ('<break time="2s"/>', ""),
        "üòê": ("<emphasis level=\"reduced\">", "</emphasis>"),
        "üôÇ": ("<emphasis level=\"moderate\">", "</emphasis>"),
        "üòÅ": ("<emphasis level=\"strong\">", "</emphasis>"),
        "üîà": ("<prosody volume=\"silent\">", "</prosody>"),
        "üîâ": ("<prosody volume=\"medium\">", "</prosody>"),
        "üîä": ("<prosody volume=\"loud\">", "</prosody>"),
        "üåè": ("<lang xml:lang=\"en-US\">", "</lang>")
    }
    supported_text_files = ["txt", "md", "rst", "docx"]

    def __init__(self, title: str, **kwargs):
        super(MainScreen, self).__init__(**kwargs)
        self.title = title
        self.ids.text_main.font_size = 24
        self.original_font_size = self.ids.text_main.font_size
        # Initialize the current TTS engine text
        self.update_current_engine_text()
        self.last_path = None
        self.opened_file = None
        # FIXME This is used to keep track of the file manager state (open or closed) but is not currently used
        self.manager_open = False
        self.file_manager = MDFileManager(
            exit_manager=self.exit_manager,
            select_path=self.select_path,
            icon_selection_button="folder-marker"
        )
        # TODO Adjust the scrollbar within MDFileManager to be more visible (not just a thin line)

        # new: for cursor movement and selection of words
        self.cnt_button = 0
        self.old_cnt_button = 0
        self.old_cursor_index = self.ids.text_main.cursor_index()
        Clock.schedule_once(self.set_focus, 0.1)
        self.load_current_voice()
        api_factory.get_active_api().settings.bind(voice_text=self.update_current_voice)

    def change_text_size(self):
        try:
            # Access the current font size
            current_font_size = self.ids.text_main.font_size

            # Define the maximum and original font size
            max_font_size = 48

            # Toggle between increasing and resetting the font size
            if current_font_size < max_font_size:
                # Increase the font size
                new_font_size = min(max_font_size, current_font_size + 2)
            else:
                # Reset to the original font size
                new_font_size = self.original_font_size

            # Apply the new font size
            self.ids.text_main.font_size = new_font_size

            # Log the change
            log.info(f"Text size changed to {new_font_size}")
        except Exception as e:
            log.error(f"Error changing text size: {e}")

    def on_ssml_button_click(self):
        try:
            log.debug("%s: SSML button pressed", self.__class__.__name__)
            emoji_popup = EmojiPopup(self.ids.text_main)
            emoji_popup.open()
        except Exception as e:
            log.error("%s: Error with SSML button: %s", self.__class__.__name__, e)

    def load_current_voice(self):
        app_instance = App.get_running_app()
        current_engine = self.get_current_tts_engine()

        try:
            # Retrieve the current voice for the selected engine
            self.selected_voice = api_factory.get_active_api().get_voice_name()
        except:
            log.error(f"Could not retrieve set voice name of current TTS engine {current_engine}")
            self.selected_voice = ""

        # Update the current voice display
        self.ids.btn_select_voice.text = f"current voice:\n{self.selected_voice}"

    def update_current_voice(self, instance, value):
        self.selected_voice = value if value is not None else ""
        self.ids.btn_select_voice.text = f"current voice:\n{self.selected_voice}"

    def on_menu_open(self):
        menu_items = [
            {
                "text": option,
                "on_release": lambda x=option: self.menu_callback(x),
            } for option in self.menu_options.keys()
        ]
        self.drop_menu = MDDropdownMenu(
            caller=self.ids.btn_menu, items=menu_items
        )
        self.drop_menu.open()

    def menu_callback(self, text_item):
        self.manager.transition.direction = 'left'
        if text_item not in self.menu_options.keys():
            log.error("%s: Invalid menu option: %s",
                      self.__class__.__name__, text_item)
            return
        if text_item == "Exit":
            ExitDialog().open()
        self.manager.current = self.menu_options[text_item]
        self.drop_menu.dismiss()

    def select_path(self, path):
        log.info("%s: Selected path: %s", self.__class__.__name__, path)
        if os.path.isfile(path):
            self.opened_file = os.path.basename(path)
            self.last_path = os.path.dirname(path)
            log.debug("%s: File: %s - Path: %s", self.__class__.__name__,
                      self.opened_file, self.last_path)
        elif os.path.isdir(path):
            self.last_path = path
        else:
            log.error("%s: Invalid path selected: %s",
                      self.__class__.__name__, path)
        self.exit_manager()

    def exit_manager(self, *args):
        if all([self.last_path, self.opened_file]):
            file = os.path.join(self.last_path, self.opened_file)
            self.load_text_from_file(file)
        else:
            log.error("%s: No file selected. Last path: %s",
                      self.__class__.__name__, self.last_path)
        self.manager_open = False
        self.file_manager.close()

    def docx_to_text(self, file_path):
        text = ""
        try:
            text = docx2txt.process(file_path)
        except Exception as e:
            log.error("%s: Error reading DOCX file: %s. Exception: %s",
                      self.__class__.__name__, file_path, e)
        return text

    def save_docx_text(self, file_path, text):
        try:
            # Open the existing document
            document = Document(file_path)

            # Overwrite content like in text file
            for _ in range(len(document.paragraphs)):
                p = document.paragraphs[0]
                p._element.getparent().remove(p._element)

            # Write text to document
            for line in text.split('\n'):
                document.add_paragraph(line)

            # Save document
            document.save(file_path)
            log.info("%s: Saved DOCX file: %s", self.__class__.__name__, file_path)
        except Exception as e:
            log.error("%s: Error saving DOCX file: %s. Exception: %s",
                      self.__class__.__name__, file_path, e)

    def on_load_file(self):
        if self.last_path is not None:
            path = self.last_path
        else:
            path = os.path.expanduser("~")
        self.file_manager.show(path)
        self.manager_open = True

    def on_save_file(self):
        if self.last_path is None or self.opened_file is None:
            log.error("%s: No file opened to save.", self.__class__.__name__)
            return
        file = os.path.join(self.last_path, self.opened_file)
        file_ext = os.path.splitext(file)[1][1:].lower()

        if file_ext == "docx":
            self.save_docx_text(file, self.ids.text_main.text)
        elif file_ext == "pdf":
            log.info("%s: Saving Pdfs not supported yet.", self.__class__.__name__)
        else:
            self.save_textfile(file)

    def load_text_from_file(self, file: str):
        if file is None:
            log.error("%s: No file selected to load.", self.__class__.__name__)
            return
        if not os.path.isfile(file):
            log.error("%s: Selection is not a file: %s", self.__class__.__name__, file)
            return
        file_ext = os.path.splitext(file)[1][1:].lower()

        if file_ext not in self.supported_text_files:
            log.error("%s: Unsupported file type: %s. Supported types: %s",
                      self.__class__.__name__, file_ext, self.supported_text_files)
            self.opened_file = None
            self.ids.text_main.text = ""
            return

        try:
            if file_ext == "docx":
                text = self.docx_to_text(file)
            else:  # For txt, md, rst
                try:
                    with open(os.path.abspath(file), 'r', encoding='utf-8') as f:
                        text = f.read()
                except UnicodeDecodeError:
                    log.warning("%s: UTF-8 decoding failed. Trying ISO-8859-1.", self.__class__.__name__)
                    with open(os.path.abspath(file), 'r', encoding='iso-8859-1') as f:
                        text = f.read()

            self.ids.text_main.text = text
            log.info("%s: Loaded file: %s", self.__class__.__name__, file)
            log.debug("%s: Text: %s...", self.__class__.__name__, text[0:40])
        except Exception as e:
            log.error("%s: Error loading file: %s. Exception: %s", self.__class__.__name__, file, e)

    def save_textfile(self, file: str):
        if file is None:
            log.error("%s: No file selected to save.", self.__class__.__name__)
            return
        with open(os.path.abspath(file), 'w') as file:
            file.write(self.ids.text_main.text)
            log.info("%s: Saved file: %s", self.__class__.__name__, file)

    def get_current_tts_engine(self):
        # Retrieve the current TTS engine from global settings
        app = App.get_running_app()
        current_engine = app.global_settings.get_setting("TTS", "current_engine",
                                                         default=api_factory.get_default_api_name())
        api_factory.set_active_api_name(current_engine)
        return current_engine

    def update_current_engine_text(self):
        # Update the property with the current engine text
        self.current_engine_text = f"tts engine: \n{self.get_current_tts_engine()}"

    def on_select_tts_engine(self):
        api = api_factory.get_active_api()
        available_engines = api_factory.get_apis_dict()

        # Prepare menu items for each available TTS engine
        menu_items = [
            {
                "text": engine_name,
                "on_release": lambda x=engine_name: self.select_tts_engine(x),
            } for engine_name in available_engines.keys()
        ]

        # Create and open the dropdown menu
        self.engine_dropdown_menu = MDDropdownMenu(
            caller=self.ids.btn_select_engine,
            items=menu_items,
            width_mult=4
        )
        self.engine_dropdown_menu.open()

    def select_tts_engine(self, engine_name):
        log.info("Selected TTS engine: %s", engine_name)

        api_factory.set_active_api_name(engine_name)
        self.ids.btn_select_engine.text = f"tts engine:\n{engine_name}"

        # Retrieve the current voice for the selected engine
        try:
            self.selected_voice = api_factory.get_active_api().get_voice_name()
        except Exception:
            log.error("Getting current voice of selected API failed.")
            self.selected_voice = ""

        # Update the selected TTS engine in global settings or any relevant setting location
        App.get_running_app().global_settings.update_setting("TTS", "current_engine", engine_name)

        # Update the button text to display the selected engine
        self.update_current_engine_text()

        # Update the current voice display
        self.update_current_voice(None, self.selected_voice)

        # Dismiss the dropdown menu
        self.engine_dropdown_menu.dismiss()

    def on_select_voice(self):
        api = api_factory.get_active_api()
        current_engine = self.get_current_tts_engine()

        if api:
            voice_names = api.get_available_voice_names()
            if voice_names:
                # Create dropdown menu items for each available voice
                menu_items = [
                    {
                        "text": voice_name,  # Display the formatted voice name
                        "on_release": lambda x=voice_name: self.select_voice(x),
                    } for voice_name in voice_names
                ]
                # Create and open the dropdown menu for voice selection
                self.voice_dropdown_menu = MDDropdownMenu(
                    caller=self.ids.btn_select_voice,
                    items=menu_items,
                    width_mult=4,
                )
                self.voice_dropdown_menu.open()
            else:
                log.error("%s: No voices available for the selected TTS engine.", self.__class__.__name__)
        else:
            log.error("%s: API not available for the selected TTS engine.", self.__class__.__name__)

    def select_voice(self, voice_name):
        # process selected voice names
        log.info("%s: Selected voice: %s", self.__class__.__name__, voice_name)
        api = api_factory.get_active_api()
        api.set_voice_name(voice_name)

        popup_window = CustomPopup(content_text=f"You selected the voice: \n{voice_name}",
                                   size_hint=(None, None), size=(400, 400))
        popup_window.open()

        # Update the current voice display
        self.update_current_voice(None, voice_name)

        self.voice_dropdown_menu.dismiss()

    def on_play(self):
        # TODO Implement audio playback (this is mostly a placeholder without a backend implementation yet)
        api = api_factory.get_active_api()
        if api:
            try:
                api.play()
            except NotImplementedError:
                log.error(
                    "%s: Audio playback not implemented for this API.", self.__class__.__name__)
            except Exception as e:
                log.error("%s: Error during playback: %s",
                          self.__class__.__name__, e)

    def on_synthesize(self):
        # TODO Implement text to speech synthesis (this is mostly a placeholder without a backend implementation yet)
        api = api_factory.get_active_api()
        tmp_dir = App.get_running_app().global_settings.get_tmp_dir()
        synthesized_file = os.path.join(tmp_dir, 'output_file.wav')
        log.info(f"Using synthesized_file={synthesized_file}")

        if api:
            msg = f"Text has been synthesized\nto an audio file"
            try:
                # FIXME: Use constant or configurable output path
                api.synthesize(self.ids.text_main.text, synthesized_file)
            except NotImplementedError:
                msg = "Text to speech synthesis not implemented for this API."
                log.error("%s: %s", self.__class__.__name__, msg)

            except Exception as e:
                msg = "Error during synthesis: \n" + str(e)
                log.error("%s: %s: %s", self.__class__.__name__, msg, e)

            self.ids.label_status.text = msg
            popup_window = CustomPopup(content_text=msg,
                                       size_hint=(None, None), size=(400, 400))
            popup_window.open()

    def on_cursor_control(self):
        new_cursor_index = self.ids.text_main.cursor_index()
        old_cursor_index = self.old_cursor_index
        cnt_button = self.cnt_button
        old_cnt_button = self.old_cnt_button

        if abs(cnt_button - old_cnt_button) == 0 and abs(new_cursor_index - old_cursor_index) > 1:
            # set cursor to the end of the word
            cursor_in_row_index = self.ids.text_main._cursor[0]
            row = self.ids.text_main._cursor[1]
            text_row = self.ids.text_main._lines[row]

            spaces = []
            for index in range(cursor_in_row_index, len(text_row)):
                char = text_row[index]
                if char == " ":
                    spaces.append(index)
            if spaces:
                end_word_index = spaces[0]
                self.ids.text_main.cursor = (end_word_index, row)
            else:
                self.ids.text_main.do_cursor_movement('cursor_end')

        self.new_cursor_index = self.ids.text_main.cursor_index()
        self.save_old_cursor_index(new_cursor_index)
        self.save_old_cnt_button(cnt_button)

    def save_old_cursor_index(self, cursor_index):
        self.old_cursor_index = cursor_index

    def save_old_cnt_button(self, cnt_button):
        self.old_cnt_button = cnt_button

    def on_cursor_left(self):
        self.ids.text_main.do_cursor_movement('cursor_left')
        self.cnt_button = self.cnt_button + 1
        Clock.schedule_once(self.set_focus, 0.1)

    def on_cursor_right(self):
        self.ids.text_main.do_cursor_movement('cursor_right')
        self.cnt_button = self.cnt_button + 1
        Clock.schedule_once(self.set_focus, 0.1)

    def set_focus(self, dt):
        self.ids.text_main.focus = True


class CustomPopup(Popup):
    content_text = StringProperty("")

    def __init__(self, **kwargs):
        super(CustomPopup, self).__init__(**kwargs)


class Popups(FloatLayout):
    voice_name = StringProperty("")

    def __init__(self, voice_name="", **kwargs):
        super(Popups, self).__init__(**kwargs)
        self.voice_name = voice_name
